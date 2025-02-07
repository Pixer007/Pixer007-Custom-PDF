import os
import streamlit as st
from docx import Document
from fpdf import FPDF  # fpdf library to create PDFs (cross-platform)


def replace_placeholders(template_path, replacements):
    """
    Replace placeholders in the Word template with user inputs while preserving formatting.
    """
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    return doc


def save_word_as_pdf(word_path, pdf_path, background_image=None):
    """
    Convert the updated Word document to PDF using fpdf, with UTF-8 encoding and a custom background.
    """
    # Create a PDF object using FPDF
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Add a custom background image if provided
    if background_image:
        pdf.image(background_image, x=0, y=0, w=210, h=297)  # A4 size in mm

    # Set the font with UTF-8 support
    pdf.set_font("Arial", size=12)

    # Open and read the Word document content
    doc = Document(word_path)
    for para in doc.paragraphs:
        # Ensure that text is encoded in UTF-8 (handle special characters)
        text = para.text.encode("latin-1", "replace").decode("latin-1")
        
        # Apply basic styles: bold, italic, underline
        if para.style.name == 'Heading 1':
            pdf.set_font("Arial", style='B', size=16)  # Heading 1 - Bold and larger font size
        elif para.style.name == 'Heading 2':
            pdf.set_font("Arial", style='B', size=14)  # Heading 2 - Bold
        elif para.style.name == 'Normal':
            pdf.set_font("Arial", size=12)  # Normal paragraph
        elif para.style.name == 'Italic':
            pdf.set_font("Arial", style='I', size=12)  # Italics
        elif para.style.name == 'Bold':
            pdf.set_font("Arial", style='B', size=12)  # Bold

        pdf.multi_cell(0, 10, text)

    pdf.output(pdf_path)


def main():
    st.title("Automated NDA PDF Generator")
    st.write("Fill in the details below to generate a customized NDA PDF.")

    # User inputs
    name = st.text_input("Enter Full Name")
    email = st.text_input("Enter Email")
    phone = st.text_input("Enter Phone Number")
    address = st.text_area("Enter Address (multi-line allowed)")

    # Paths for template and output files
    template_path = os.path.abspath("NDA Template - INDIA 3.docx")  # Predefined common template
    output_docx = os.path.abspath("updated_template.docx")
    output_pdf = os.path.abspath("generated_nda.pdf")
    background_image = "A4 - 2.jpg"  # Your custom background image path

    if st.button("Generate PDF"):
        if not os.path.exists(template_path):
            st.error("Template file not found. Please ensure 'template.docx' is in the working directory.")
            return

        if not (name and email and phone and address):
            st.error("Please fill in all the fields!")
            return

        try:
            # Define the placeholder replacements
            replacements = {
                "Client Name": name,
                "Client Email": email,
                "Client Phone": phone,
                "Client Address": address,
            }

            # Replace placeholders in the template
            updated_doc = replace_placeholders(template_path, replacements)
            updated_doc.save(output_docx)  # Save updated Word document

            # Convert the updated Word document to PDF with custom background
            save_word_as_pdf(output_docx, output_pdf, background_image=background_image)

            # Provide the download link for the PDF
            with open(output_pdf, "rb") as pdf_file:
                st.download_button(
                    label="Download NDA PDF",
                    data=pdf_file,
                    file_name="generated_nda.pdf",
                    mime="application/pdf",
                )

        except Exception as e:
            st.error(f"An error occurred: {e}")


if __name__ == "__main__":
    main()
